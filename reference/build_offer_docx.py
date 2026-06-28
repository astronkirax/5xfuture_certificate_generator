# -*- coding: utf-8 -*-
"""Builds a portrait A4 internship OFFER LETTER (.docx) for 5XFUTURE with fillable placeholders.

Matches the branding of internship-certificate.html / build_docx.py:
clean corporate letterhead, real contact details, online/remote unpaid training
internship, the Director's signature (signature.png) over the sign-off, a student
acceptance block, and a space to affix the office seal. Tuned to one A4 page.
"""
import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Brand palette (sampled from the 5XFUTURE logo) ----
NAVY        = "3A266D"   # brand primary  (deep purple)
NAVY_SOFT   = "4A3488"
GOLD        = "F58220"   # brand accent   (orange)
GOLD_DEEP   = "D96E12"
MUTED       = "6E6688"
INK         = "2B2347"
CREAM       = "FBF8F2"
PLACEHOLDER = "C3BDD2"
WHITE       = "EFEAF7"

SERIF_HEAD = "Cambria"            # elegant heading serif (ships with Office/Windows)
SERIF_BODY = "Palatino Linotype"  # refined body serif

HERE = os.path.dirname(os.path.abspath(__file__))

# ---- Real institute contact details ----
WEBSITE = "5xfuture.in"
EMAIL   = "contact@5xfuture.in"
PHONE   = "+91 96401 31555"
CIN     = "U62010AP2023PTC113188"


# ---------- low-level helpers (shared pattern with build_docx.py) ----------
def add_run(p, text, *, name=SERIF_BODY, size=12, color=INK,
            bold=False, italic=False, underline=False, spacing=None,
            caps=False, small_caps=False):
    r = p.add_run(text)
    r.font.name = name
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), name)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.underline = underline
    r.font.color.rgb = RGBColor.from_string(color)
    if caps:
        e = OxmlElement('w:caps'); e.set(qn('w:val'), 'true'); rpr.append(e)
    if small_caps:
        e = OxmlElement('w:smallCaps'); e.set(qn('w:val'), 'true'); rpr.append(e)
    if spacing is not None:                      # spacing in points -> twips (x20)
        e = OxmlElement('w:spacing'); e.set(qn('w:val'), str(int(spacing * 20))); rpr.append(e)
    return r


def blank(p, n=14, size=12):
    """An underlined fill-in blank (placeholder grey)."""
    return add_run(p, " " * n, size=size, color=PLACEHOLDER, underline=True)


def set_space(p, before=0, after=0, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line
    return p


def para_bottom_border(p, color=NAVY, sz=8, space=2):
    """Draw an underline rule beneath a paragraph (signature / name lines)."""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), str(space))
    b.set(qn('w:color'), color)
    pbdr.append(b)
    pPr.append(pbdr)


def set_left_page_border(section, color=NAVY, sz=24, space=18):
    """A single thick accent rule down the left page edge (clean letterhead look)."""
    sectPr = section._sectPr
    pgB = OxmlElement('w:pgBorders')
    pgB.set(qn('w:offsetFrom'), 'page')
    e = OxmlElement('w:left')
    e.set(qn('w:val'), 'single')
    e.set(qn('w:sz'), str(sz))
    e.set(qn('w:space'), str(space))
    e.set(qn('w:color'), color)
    pgB.append(e)
    sectPr.append(pgB)


def set_cell_border(cell, color=GOLD, sz=6, val="single"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), val)
        e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    tcPr.append(borders)


def set_cell_left_accent(cell, color=GOLD, sz=24):
    """Thick left border only — for the acceptance call-out box."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    e = OxmlElement('w:left')
    e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
    e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
    borders.append(e)
    for edge in ('top', 'bottom', 'right'):
        b = OxmlElement('w:' + edge)
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0'); b.set(qn('w:color'), NAVY)
        borders.append(b)
    tcPr.append(borders)


def set_cell_margins(cell, top=120, bottom=120, left=200, right=200):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for edge, v in (('top', top), ('bottom', bottom), ('start', left),
                    ('left', left), ('end', right), ('right', right)):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:w'), str(v))
        e.set(qn('w:type'), 'dxa')
        m.append(e)
    tcPr.append(m)


def cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'none')
        borders.append(e)
    tblPr.append(borders)


def set_row_height(row, mm, exact=True):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight')
    h.set(qn('w:val'), str(int(Mm(mm).twips)))
    h.set(qn('w:hRule'), 'exact' if exact else 'atLeast')
    trPr.append(h)


def clear_cell(cell):
    """Remove the default empty paragraph so we control spacing inside a cell."""
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)


def bullet(text_runs):
    """A brand-styled bullet list item. text_runs: list of (text, kwargs)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Mm(8)
    pf.first_line_indent = Mm(-4)
    set_space(p, before=0, after=2, line=1.08)
    add_run(p, "•  ", size=11, color=GOLD_DEEP, bold=True)
    for t, kw in text_runs:
        add_run(p, t, **kw)
    return p


# ---------- document ----------
doc = Document()
# Drop any default leading empty paragraph so the letterhead starts at the top.
if doc.paragraphs:
    doc.paragraphs[0]._p.getparent().remove(doc.paragraphs[0]._p)

# Portrait A4
sec = doc.sections[0]
sec.page_width = Mm(210)
sec.page_height = Mm(297)
sec.top_margin = Mm(11)
sec.bottom_margin = Mm(10)
sec.left_margin = Mm(18)
sec.right_margin = Mm(18)
set_left_page_border(sec, color=NAVY, sz=24, space=14)

CONTENT_W = Mm(210) - Mm(18) - Mm(18)   # usable width between margins
BODY = 11.5    # body font size (pt)
LINE = 1.13    # body line spacing

# Default style font
doc.styles['Normal'].font.name = SERIF_BODY
doc.styles['Normal'].font.size = Pt(BODY)
doc.styles['Normal'].paragraph_format.space_after = Pt(0)


# ===== Letterhead: logo (left) + contact (right) =====
head = doc.add_table(rows=1, cols=2)
no_table_borders(head)
head.columns[0].width = int(CONTENT_W * 0.55)
head.columns[1].width = int(CONTENT_W * 0.45)

# left: logo + sub-line
lc = head.cell(0, 0); clear_cell(lc)
lp = lc.add_paragraph(); set_space(lp, before=0, after=2)
lp.add_run().add_picture(os.path.join(HERE, "logo.png"), width=Mm(46))
sp = lc.add_paragraph(); set_space(sp, before=0, after=0)
add_run(sp, "SOFTWARE TRAINING INSTITUTE  •  KAKINADA",
        size=8.5, color=MUTED, bold=True, spacing=1.5)

# right: real contact details (right-aligned). Plain-text labels render in any Word.
rc = head.cell(0, 1); clear_cell(rc)
rc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
for pre, val in (("Web:  ", WEBSITE), ("Email:  ", EMAIL), ("Phone:  ", PHONE)):
    cp = rc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_space(cp, before=0, after=1)
    add_run(cp, pre, size=10, color=MUTED)
    add_run(cp, val, size=10, color=NAVY, bold=True)

# brand rule (thick purple underline under the letterhead)
rule = doc.add_paragraph(); set_space(rule, before=2, after=0)
para_bottom_border(rule, color=NAVY, sz=20, space=1)
add_run(rule, " ", size=2)

# ===== Title =====
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_space(p, before=6, after=5)
add_run(p, "INTERNSHIP OFFER LETTER", name=SERIF_HEAD, size=18, color=NAVY,
        bold=True, spacing=3)

# ===== Reference row: Ref No. (left) + Date (right) =====
ref = doc.add_table(rows=1, cols=2)
no_table_borders(ref)
ref.columns[0].width = int(CONTENT_W * 0.5)
ref.columns[1].width = int(CONTENT_W * 0.5)
rl = ref.cell(0, 0); clear_cell(rl)
rp = rl.add_paragraph(); set_space(rp, before=2, after=0)
add_run(rp, "Ref / Letter No.: ", size=BODY, color=INK); blank(rp, 12, size=BODY)
rr = ref.cell(0, 1); clear_cell(rr)
rp = rr.add_paragraph(); rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_space(rp, before=2, after=0)
add_run(rp, "Date: ", size=BODY, color=INK); blank(rp, 14, size=BODY)

# ===== Recipient =====
p = doc.add_paragraph(); set_space(p, before=8, after=0)
add_run(p, "To,", size=BODY, color=INK)
p = doc.add_paragraph(); set_space(p, before=3, after=0)
blank(p, 30, size=BODY); add_run(p, "   (Name of the Student)", size=9, color=MUTED)
p = doc.add_paragraph(); set_space(p, before=3, after=0)
blank(p, 30, size=BODY); add_run(p, "   (City / Place - optional)", size=9, color=MUTED)

# ===== Subject =====
p = doc.add_paragraph(); set_space(p, before=8, after=0)
add_run(p, "Subject: ", size=BODY, color=INK, bold=True)
add_run(p, "Offer of Internship in ", size=BODY, color=NAVY, bold=True)
blank(p, 22, size=BODY)

# ===== Salutation =====
p = doc.add_paragraph(); set_space(p, before=7, after=0)
add_run(p, "Dear ", size=BODY, color=INK); blank(p, 16, size=BODY)
add_run(p, ",", size=BODY, color=INK)

# ===== Body =====
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
set_space(p, before=7, after=0, line=LINE)
add_run(p, "We are pleased to offer you an internship at ", size=BODY, color=INK)
add_run(p, "5XFUTURE Software Training Institute, Kakinada", size=BODY, color=NAVY, bold=True)
add_run(p, " - a brand of ", size=BODY, color=INK)
add_run(p, "Mahasana IT Solutions Private Limited", size=BODY, color=NAVY, bold=True)
add_run(p, " - in the domain of ", size=BODY, color=INK)
blank(p, 18, size=BODY)
add_run(p, ". The internship will be conducted in ", size=BODY, color=INK)
add_run(p, "online / offline", size=BODY, color=NAVY, bold=True)
add_run(p, " mode.", size=BODY, color=INK)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
set_space(p, before=6, after=0, line=LINE)
add_run(p, "Your internship will commence on ", size=BODY, color=INK)
blank(p, 12, size=BODY)
add_run(p, " and conclude on ", size=BODY, color=INK)
blank(p, 12, size=BODY)
add_run(p, ", spanning a period of ", size=BODY, color=INK)
blank(p, 8, size=BODY)
add_run(p, ". During this time you will work on ", size=BODY, color=INK)
add_run(p, "real-time projects", size=BODY, color=NAVY, bold=True)
add_run(p, " and develop ", size=BODY, color=INK)
add_run(p, "industry-relevant skills", size=BODY, color=NAVY, bold=True)
add_run(p, " under the guidance of our mentors and trainers.", size=BODY, color=INK)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
set_space(p, before=6, after=0, line=LINE)
add_run(p, "This is a skill-development training internship offered as part of our "
           "learning programme. It is ", size=BODY, color=INK)
add_run(p, "unpaid", size=BODY, color=NAVY, bold=True)
add_run(p, ", and ", size=BODY, color=INK)
add_run(p, "no stipend is payable", size=BODY, color=NAVY, bold=True)
add_run(p, " for the duration of the internship.", size=BODY, color=INK)

# Terms
p = doc.add_paragraph(); set_space(p, before=6, after=2)
add_run(p, "Terms of the Internship:", size=BODY, color=NAVY, bold=True)
bullet([("Attend all scheduled online sessions and submit the assigned tasks and "
         "projects on time.", dict(size=11, color=INK))])
bullet([("Maintain professionalism, punctuality, discipline, and confidentiality of "
         "all training and project material.", dict(size=11, color=INK))])
bullet([("On successful completion, you will be awarded a ", dict(size=11, color=INK)),
        ("Certificate of Internship", dict(size=11, color=NAVY, bold=True)),
        (" by 5XFUTURE.", dict(size=11, color=INK))])
bullet([("The institute reserves the right to discontinue the internship in case of "
         "misconduct, irregular attendance, or non-participation.", dict(size=11, color=INK))])

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
set_space(p, before=6, after=0, line=LINE)
add_run(p, "We are confident that this internship will be a valuable step in your "
           "career journey. To confirm your participation, kindly sign and return "
           "the acceptance section below.", size=BODY, color=INK)

# ===== Closing: greeting (left) + authorised signatory with signature (right) =====
clo = doc.add_table(rows=1, cols=2)
no_table_borders(clo)
clo.columns[0].width = int(CONTENT_W * 0.40)
clo.columns[1].width = int(CONTENT_W * 0.60)

sl = clo.cell(0, 0); clear_cell(sl); sl.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
rp = sl.add_paragraph(); set_space(rp, before=8, after=0)
add_run(rp, "Warm regards,", size=BODY, color=INK)

# right: the Director's signature is the official sign-off (no empty seal box)
sr = clo.cell(0, 1); clear_cell(sr); sr.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
ip = sr.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_space(ip, before=4, after=0)
ip.add_run().add_picture(os.path.join(HERE, "signature.png"), width=Mm(52))
cap = sr.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_space(cap, before=0, after=0)
add_run(cap, "Authorised Signatory", name=SERIF_HEAD, size=10, color=NAVY, bold=True, spacing=1)

# ===== Acceptance call-out =====
acc = doc.add_table(rows=1, cols=1)
acc.alignment = WD_TABLE_ALIGNMENT.CENTER
ac = acc.cell(0, 0); ac.width = CONTENT_W
set_cell_left_accent(ac, color=GOLD, sz=24)
cell_shading(ac, CREAM)
set_cell_margins(ac, top=90, bottom=100, left=240, right=240)
clear_cell(ac)
hp = ac.add_paragraph(); set_space(hp, before=0, after=3)
add_run(hp, "ACCEPTANCE BY INTERN", name=SERIF_HEAD, size=10.5, color=NAVY, bold=True, spacing=2)
tp = ac.add_paragraph(); set_space(tp, before=0, after=5, line=1.1)
add_run(tp, "I, ", size=11, color=INK); blank(tp, 22, size=11)
add_run(tp, ", hereby accept the internship offer and agree to the terms and "
            "conditions stated above.", size=11, color=INK)
sp = ac.add_paragraph(); set_space(sp, before=0, after=0)
add_run(sp, "Student Signature: ", size=11, color=INK); blank(sp, 16, size=11)
add_run(sp, "        Date: ", size=11, color=INK); blank(sp, 12, size=11)

# ===== Footer band (full-width purple) =====
foot = doc.add_table(rows=1, cols=1)
foot.alignment = WD_TABLE_ALIGNMENT.CENTER
fc = foot.cell(0, 0); fc.width = CONTENT_W
cell_shading(fc, NAVY)
set_cell_margins(fc, top=60, bottom=60, left=120, right=120)
clear_cell(fc)
fp = fc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_space(fp, 0, 0)
add_run(fp, "5XFUTURE is a brand of ", size=8.5, color=WHITE)
add_run(fp, "Mahasana IT Solutions Private Limited", size=8.5, color=GOLD, bold=True)
add_run(fp, " (CIN: " + CIN + ")   •   " + EMAIL + "   •   " + PHONE, size=8.5, color=WHITE)

# Word keeps a paragraph after the final table; minimise it so it never spills
# onto a second page when page 1 is full.
tail = doc.add_paragraph()
tail.paragraph_format.space_before = Pt(0)
tail.paragraph_format.space_after = Pt(0)
tail.paragraph_format.line_spacing = Pt(1)
_pPr = tail._p.get_or_add_pPr()
_rPr = OxmlElement('w:rPr'); _sz = OxmlElement('w:sz'); _sz.set(qn('w:val'), '2')
_rPr.append(_sz); _pPr.append(_rPr)

out = os.path.join(HERE, "offer-letter.docx")
doc.save(out)
print("Saved:", out)
