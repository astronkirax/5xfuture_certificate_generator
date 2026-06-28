# -*- coding: utf-8 -*-
"""Builds a landscape internship certificate (.docx) for 5XFUTURE with fillable placeholders."""
import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Brand palette (sampled from the 5XFUTURE logo) ----
NAVY        = "3A266D"   # brand primary  (deep purple)
NAVY_SOFT   = "4A3488"
GOLD        = "F58220"   # brand accent   (orange)
GOLD_DEEP   = "D96E12"
MUTED       = "6E6688"
PLACEHOLDER = "C3BDD2"

SERIF_HEAD = "Cambria"          # elegant heading serif (ships with Office/Windows)
SERIF_BODY = "Palatino Linotype"  # refined body serif


# ---------- low-level helpers ----------
def add_run(p, text, *, name=SERIF_BODY, size=12, color=NAVY,
            bold=False, italic=False, underline=False, spacing=None,
            caps=False, small_caps=False):
    r = p.add_run(text)
    r.font.name = name
    # ensure complex/east-asian scripts use same font
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), name)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.underline = underline
    r.font.color.rgb = RGBColor.from_string(color)
    if caps:
        e = OxmlElement('w:caps'); e.set(qn('w:val'), 'true'); rpr.append(e)
    if small_caps:
        e = OxmlElement('w:smallCaps'); e.set(qn('w:val'), 'true'); rpr.append(e)
    if spacing is not None:                      # spacing in points -> twips (x20)
        e = OxmlElement('w:spacing'); e.set(qn('w:val'), str(int(spacing * 20))); rpr.append(e)
    return r


def set_space(p, before=0, after=0, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line
    return p


def para_bottom_border(p, color=NAVY, sz=8, space=2):
    """Draw an underline rule beneath a paragraph (signature / name lines)."""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), str(space))
    b.set(qn('w:color'), color)
    pbdr.append(b)
    pPr.append(pbdr)


def set_page_border(section, color=NAVY, sz=20, val="double", space=24):
    sectPr = section._sectPr
    pgB = OxmlElement('w:pgBorders')
    pgB.set(qn('w:offsetFrom'), 'page')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), val)
        e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), str(space))
        e.set(qn('w:color'), color)
        pgB.append(e)
    sectPr.append(pgB)


def set_cell_border(cell, color=GOLD, sz=6, val="single"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), val)
        e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    tcPr.append(borders)


def set_cell_margins(cell, top=120, bottom=120, left=200, right=200):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for edge, v in (('top', top), ('bottom', bottom), ('start', left),
                    ('left', left), ('end', right), ('right', right)):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:w'), str(v))
        e.set(qn('w:type'), 'dxa')
        m.append(e)
    tcPr.append(m)


def cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'none')
        borders.append(e)
    tblPr.append(borders)


# ---------- document ----------
doc = Document()
if doc.paragraphs:                       # drop default leading paragraph
    doc.paragraphs[0]._p.getparent().remove(doc.paragraphs[0]._p)

# Landscape A4
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width = Mm(297)
sec.page_height = Mm(210)
sec.top_margin = Mm(8)
sec.bottom_margin = Mm(8)
sec.left_margin = Mm(13)
sec.right_margin = Mm(13)
set_page_border(sec, color=NAVY, sz=22, val="double", space=22)

# Default style font
doc.styles['Normal'].font.name = SERIF_BODY
doc.styles['Normal'].font.size = Pt(12)

# ===== Inner gold frame = single-cell table holding all content =====
content_w = Mm(297) - Mm(13) - Mm(13)   # page width minus L/R margins
frame = doc.add_table(rows=1, cols=1)
frame.alignment = WD_TABLE_ALIGNMENT.CENTER
frame.autofit = False
cell = frame.cell(0, 0)
cell.width = content_w
cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
set_cell_border(cell, color=GOLD, sz=6, val="single")
set_cell_margins(cell, top=160, bottom=160, left=360, right=360)

# remove the default empty paragraph in the cell so we control spacing
cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)


def cpara(align=WD_ALIGN_PARAGRAPH.CENTER):
    p = cell.add_paragraph()
    p.alignment = align
    return p


# --- Brand header (logo image) ---
_logo = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo.png")
p = cpara(); set_space(p, before=2, after=0)
p.add_run().add_picture(_logo, width=Mm(66))

p = cpara(); set_space(p, before=4, after=2)
add_run(p, "S O F T W A R E   T R A I N I N G   I N S T I T U T E   •   K A K I N A D A",
        name=SERIF_BODY, size=10, color=MUTED, bold=True, spacing=0.5)

# divider
p = cpara(); set_space(p, before=2, after=2)
add_run(p, "✦", name=SERIF_BODY, size=11, color=GOLD_DEEP)

# --- Title ---
p = cpara(); set_space(p, before=4, after=0)
add_run(p, "CERTIFICATE OF INTERNSHIP", name=SERIF_HEAD, size=30, color=NAVY,
        bold=True, spacing=3)
p = cpara(); set_space(p, before=1, after=6)
add_run(p, "OF SUCCESSFUL COMPLETION", name=SERIF_BODY, size=12, color=GOLD_DEEP,
        bold=True, spacing=3)

# --- Recipient ---
p = cpara(); set_space(p, before=4, after=2)
add_run(p, "This certificate is proudly presented to", name=SERIF_BODY, size=13,
        color=MUTED, italic=True)

# Name line — blank underline to be filled in (no placeholder text)
p = cpara(); set_space(p, before=2, after=0)
para_bottom_border(p, color=NAVY, sz=8, space=3)
add_run(p, " ", name=SERIF_HEAD, size=24, color=PLACEHOLDER, bold=True)

p = cpara(); set_space(p, before=2, after=4)
add_run(p, "NAME OF THE INTERN", name=SERIF_BODY, size=9, color=MUTED, spacing=1.5)

# --- Body statement (justified, with inline blanks) ---
p = cpara(WD_ALIGN_PARAGRAPH.CENTER); set_space(p, before=4, after=2, line=1.5)
add_run(p, "for successfully completing the internship programme in  ",
        name=SERIF_BODY, size=12.5, color=NAVY)
add_run(p, " " * 18, name=SERIF_BODY, size=12.5,
        color=PLACEHOLDER, underline=True)
add_run(p, "  at ", name=SERIF_BODY, size=12.5, color=NAVY)
add_run(p, "5XFUTURE Software Training Institute, Kakinada", name=SERIF_BODY,
        size=12.5, color=NAVY, bold=True)
add_run(p, ", from ", name=SERIF_BODY, size=12.5, color=NAVY)
add_run(p, " " * 12, name=SERIF_BODY, size=12.5, color=PLACEHOLDER,
        underline=True)
add_run(p, " to ", name=SERIF_BODY, size=12.5, color=NAVY)
add_run(p, " " * 12, name=SERIF_BODY, size=12.5, color=PLACEHOLDER,
        underline=True)
add_run(p, ".", name=SERIF_BODY, size=12.5, color=NAVY)

p = cpara(WD_ALIGN_PARAGRAPH.CENTER); set_space(p, before=0, after=4, line=1.5)
add_run(p, "During the internship, the intern worked on real-time projects and "
           "demonstrated exemplary technical skill, discipline, and professionalism. "
           "We commend this achievement and wish the intern continued success ahead.",
        name=SERIF_BODY, size=12.5, color=NAVY)

# --- Signature / seal row (3-col borderless table) ---
sig = cell.add_table(rows=1, cols=3)
sig.alignment = WD_TABLE_ALIGNMENT.CENTER
no_table_borders(sig)
total = content_w - Mm(20)
sig.columns[0].width = int(total * 0.36)
sig.columns[1].width = int(total * 0.28)
sig.columns[2].width = int(total * 0.36)

def sign_cell(c, role, org, sig_img=None):
    c.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    c.paragraphs[0]._p.getparent().remove(c.paragraphs[0]._p)
    line = c.add_paragraph(); line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_space(line, before=4 if sig_img else 10, after=4)
    para_bottom_border(line, color=NAVY, sz=8, space=2)
    if sig_img:                                   # signature resting on the line
        line.add_run().add_picture(sig_img, width=Mm(34))
    else:
        add_run(line, " ", size=11)
    if role:                                      # printed role (skipped when the
        r = c.add_paragraph(); r.alignment = WD_ALIGN_PARAGRAPH.CENTER   # signature
        set_space(r, before=0, after=0)                                  # carries it
        add_run(r, role, name=SERIF_HEAD, size=11.5, color=NAVY, bold=True, spacing=1)
    o = c.add_paragraph(); o.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_space(o, before=0, after=0)
    add_run(o, org, name=SERIF_BODY, size=9, color=MUTED)

sign_cell(sig.cell(0, 0), "Training Head", "5XFUTURE, Kakinada")

# middle seal cell
mid = sig.cell(0, 1)
mid.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
mid.paragraphs[0]._p.getparent().remove(mid.paragraphs[0]._p)
seal_tbl = mid.add_table(rows=1, cols=1) if False else None  # keep simple
sp = mid.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_space(sp, before=2, after=0)
add_run(sp, "★", name=SERIF_BODY, size=16, color=GOLD)
sp2 = mid.add_paragraph(); sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_space(sp2, before=0, after=0)
add_run(sp2, "5XFUTURE", name=SERIF_HEAD, size=12, color=GOLD_DEEP, bold=True, spacing=1)
sp3 = mid.add_paragraph(); sp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_space(sp3, before=0, after=0)
add_run(sp3, "CERTIFIED", name=SERIF_BODY, size=9, color=MUTED, spacing=2)

sign_cell(sig.cell(0, 2), "", "Mahasana IT Solutions Pvt. Ltd.",
          sig_img=os.path.join(os.path.dirname(os.path.abspath(__file__)), "signature-mark.png"))

# --- Meta row: Cert No. + Date (borderless 2-col) ---
meta = cell.add_table(rows=1, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
no_table_borders(meta)
meta.columns[0].width = int((content_w - Mm(20)) * 0.5)
meta.columns[1].width = int((content_w - Mm(20)) * 0.5)

mc0 = meta.cell(0, 0)
mc0.paragraphs[0]._p.getparent().remove(mc0.paragraphs[0]._p)
mp = mc0.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_space(mp, before=8, after=0)
add_run(mp, "Certificate No. :  ", name=SERIF_BODY, size=10, color=MUTED)
add_run(mp, " " * 16, name=SERIF_BODY, size=10, color=PLACEHOLDER,
        underline=True)

mc1 = meta.cell(0, 1)
mc1.paragraphs[0]._p.getparent().remove(mc1.paragraphs[0]._p)
mp = mc1.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_space(mp, before=8, after=0)
add_run(mp, "Date of Issue :  ", name=SERIF_BODY, size=10, color=MUTED)
add_run(mp, " " * 16, name=SERIF_BODY, size=10, color=PLACEHOLDER,
        underline=True)

# --- Legal footer ---
p = cpara(); set_space(p, before=10, after=0)
add_run(p, "5XFUTURE is a brand of ", name=SERIF_BODY, size=9.5, color=MUTED)
add_run(p, "Mahasana IT Solutions Private Limited", name=SERIF_BODY, size=9.5,
        color=NAVY, bold=True)
add_run(p, " (CIN: U62010AP2023PTC113188)   |   contact@5xfuture.in   |   +91 96401 31555   |   5xfuture.in",
        name=SERIF_BODY, size=9.5, color=MUTED)

# Minimise the mandatory trailing paragraph after the frame table so the
# certificate stays on a single landscape page.
tail = doc.add_paragraph()
tail.paragraph_format.space_before = Pt(0)
tail.paragraph_format.space_after = Pt(0)
tail.paragraph_format.line_spacing = Pt(1)
_pPr = tail._p.get_or_add_pPr()
_rPr = OxmlElement('w:rPr'); _sz = OxmlElement('w:sz'); _sz.set(qn('w:val'), '2')
_rPr.append(_sz); _pPr.append(_rPr)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "internship-certificate.docx")
doc.save(out)
print("Saved:", out)
